from pathlib import Path

import pandas as pd


def parse_pdf(file_path: str) -> list[dict]:
    import pdfplumber

    parts = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                parts.append({"text": text, "page_number": i + 1})
    return parts


def parse_docx(file_path: str) -> list[dict]:
    from docx import Document

    doc = Document(file_path)
    text_parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            text_parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return [{"text": "\n".join(text_parts), "page_number": None}]


def _parse_tabular(
    file_path: str,
    file_name: str,
    sheet_name: str | None,
    columns: list[str],
    rows: list[list[str]],
) -> str:
    lines = [
        f"Konteks dokumen: {file_name}",
    ]
    if sheet_name:
        lines.append(f"Sheet: {sheet_name}")

    col_labels = [c.strip() for c in columns]

    for i, row in enumerate(rows):
        cells = []
        for j, val in enumerate(row):
            label = col_labels[j] if j < len(col_labels) else f"Kolom{j}"
            cells.append(f"{label}: {val}")
        lines.append(f"  [{i + 1}] {' | '.join(cells)}")

        if (i + 1) % 5 == 0 and (i + 1) < len(rows):
            lines.append("---")

    return "\n".join(lines)


def _detect_csv_separator(file_path: str, encoding: str = "utf-8") -> str:
    """Auto-detect CSV separator by sampling first ~2KB.

    Returns one of: ';' ',' '\\t' '|'
    """
    try:
        with open(file_path, encoding=encoding) as f:
            sample = f.read(2048)
    except Exception:
        return ","

    candidates = {";": 0, ",": 0, "\t": 0, "|": 0}
    for line in sample.splitlines()[:5]:
        for sep in candidates:
            candidates[sep] += line.count(sep)

    # Pick the separator with the highest count, default to ','
    best = max(candidates.items(), key=lambda x: x[1])
    if best[1] == 0:
        return ","
    return best[0]


def parse_csv(file_path: str) -> list[dict]:

    file_name = Path(file_path).name
    df = None
    last_error = None

    # Try multiple encodings and separators
    for encoding in ("utf-8", "latin1", "cp1252", "utf-8-sig"):
        for sep in (None, ";", ",", "\t", "|"):
            try:
                if sep is None:
                    df = pd.read_csv(
                        file_path,
                        sep=None,
                        engine="python",
                        encoding=encoding,
                        low_memory=False,
                    )
                else:
                    df = pd.read_csv(
                        file_path,
                        sep=sep,
                        encoding=encoding,
                        low_memory=False,
                    )
                # Strip whitespace from column names
                df.columns = [str(c).strip() for c in df.columns]
                last_error = None
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                last_error = e
                continue
        if df is not None:
            break

    if df is None:
        raise ValueError(
            f"Could not parse CSV {file_name}: {last_error or 'unknown error'}"
        )

    columns = df.columns.tolist()
    rows = [[str(v) for v in row.values] for _, row in df.iterrows()]
    return [{"text": _parse_tabular(file_path, file_name, None, columns, rows), "page_number": None}]


def parse_excel(file_path: str) -> list[dict]:

    file_name = Path(file_path).name
    xls = pd.ExcelFile(file_path)
    all_sheets = []
    for sheet in xls.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet)
        df = df.fillna("")
        # Strip whitespace from column names
        df.columns = [str(c).strip() for c in df.columns]
        columns = df.columns.tolist()
        rows = [[str(v) for v in row.values] for _, row in df.iterrows()]
        sheet_text = _parse_tabular(file_path, file_name, sheet, columns, rows)
        all_sheets.append(sheet_text)
    return [{"text": "\n\n".join(all_sheets), "page_number": None}]


def parse_document(file_path: str) -> list[dict]:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".csv":
        return parse_csv(file_path)
    elif ext == ".xlsx":
        return parse_excel(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
